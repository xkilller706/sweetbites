#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generar documentos DOCX limpios sin colores - Arial 12
Todos los requerimientos en tablas individuales
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_rgb):
    """Establecer color de fondo de celda"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_rgb)
    cell._element.get_or_add_tcPr().append(shading_elm)

print("Generando documento de Requerimientos Funcionales...")
print("Formato: Tablas individuales, Arial 12, sin colores")

doc = Document()

# Configurar márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Título principal
title = doc.add_heading('REQUERIMIENTOS FUNCIONALES', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True

subtitle = doc.add_paragraph('Sistema SweetBites - Gestión de Recetas de Postres')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(12)

doc.add_paragraph()

# TODOS LOS 40 REQUERIMIENTOS
requerimientos = [
    # MÓDULO 1: AUTENTICACIÓN Y USUARIOS (8)
    {
        'id': 'RF-AU-01',
        'nombre': 'Registro de Usuario',
        'modulo': 'Autenticación y Usuarios',
        'descripcion': 'El usuario nuevo podrá crear una cuenta en el sistema ingresando sus datos personales para acceder a las funcionalidades completas de la app de recetas de postres.',
        'como': 'Usuario nuevo',
        'quiero': 'Registrarme con mis datos personales en el sistema',
        'para': 'Acceder a todas las funcionalidades de la app como guardar recetas favoritas y comentar',
        'criterios': [
            'El usuario deberá ingresar su nombre completo (solo letras, sin números)',
            'El usuario deberá registrar un correo electrónico válido con dominio existente',
            'El usuario deberá crear una contraseña segura (mínimo 8 caracteres, una mayúscula y un número)',
            'El sistema deberá confirmar la contraseña antes de completar el registro',
            'El sistema deberá verificar que el correo no esté ya registrado',
            'El sistema deberá mostrar mensajes de error claros ante datos inválidos',
            'Al registrarse exitosamente, el sistema redirigirá al usuario al inicio'
        ],
        'prioridad': 'Alta',
        'notas': 'Las contraseñas se almacenan con hash bcrypt. El correo debe ser único en la base de datos.'
    },
    {
        'id': 'RF-AU-02',
        'nombre': 'Inicio de Sesión',
        'modulo': 'Autenticación y Usuarios',
        'descripcion': 'Los usuarios registrados pueden iniciar sesión con su correo electrónico y contraseña para acceder a su cuenta personal.',
        'como': 'Usuario registrado',
        'quiero': 'Iniciar sesión con mi correo y contraseña',
        'para': 'Acceder a mi cuenta y usar las funcionalidades del sistema',
        'criterios': [
            'El usuario debe ingresar un correo electrónico válido',
            'El usuario debe ingresar su contraseña',
            'El sistema debe validar las credenciales contra la base de datos',
            'Si las credenciales son correctas, se genera un token JWT válido por 7 días',
            'El token se almacena en localStorage del navegador',
            'El usuario admin es redirigido a /admin',
            'El usuario normal es redirigido a la página principal',
            'Se muestra mensaje de error si las credenciales son incorrectas'
        ],
        'prioridad': 'Alta',
        'notas': 'El token JWT incluye el ID y rol del usuario. La sesión permanece activa por 7 días.'
    },
    {
        'id': 'RF-AU-03',
        'nombre': 'Cerrar Sesión',
        'modulo': 'Autenticación y Usuarios',
        'descripcion': 'Los usuarios pueden cerrar su sesión activa de forma segura.',
        'como': 'Usuario autenticado',
        'quiero': 'Cerrar mi sesión',
        'para': 'Salir de mi cuenta de forma segura',
        'criterios': [
            'El botón de cerrar sesión debe estar visible en el menú de usuario',
            'Al hacer clic, se elimina el token del localStorage',
            'Se limpia el sessionStorage',
            'El usuario es redirigido a la página principal',
            'Se muestra un mensaje de confirmación'
        ],
        'prioridad': 'Media',
        'notas': 'La función authService.logout() maneja toda la limpieza de sesión.'
    },
    {
        'id': 'RF-AU-04',
        'nombre': 'Ver Perfil de Usuario',
        'modulo': 'Autenticación y Usuarios',
        'descripcion': 'Los usuarios pueden visualizar su información personal completa y estadísticas de actividad.',
        'como': 'Usuario autenticado',
        'quiero': 'Ver mi perfil completo',
        'para': 'Consultar mi información personal y mis estadísticas',
        'criterios': [
            'Se debe mostrar el nombre completo del usuario',
            'Se debe mostrar el correo electrónico',
            'Se debe mostrar el teléfono (si fue proporcionado)',
            'Se debe mostrar la foto de perfil o avatar por defecto',
            'Se debe mostrar la biografía del usuario',
            'Se debe mostrar el rol (usuario o admin)',
            'Se debe mostrar el plan (gratis o premium)',
            'Se debe mostrar la fecha de registro',
            'Se deben mostrar estadísticas: total recetas creadas, favoritos, comentarios'
        ],
        'prioridad': 'Alta',
        'notas': 'Página: /user/profile. Endpoint: GET /api/users/profile'
    },
    {
        'id': 'RF-AU-05',
        'nombre': 'Editar Perfil de Usuario',
        'modulo': 'Autenticación y Usuarios',
        'descripcion': 'Los usuarios pueden actualizar su información personal.',
        'como': 'Usuario autenticado',
        'quiero': 'Editar mi información personal',
        'para': 'Mantener mis datos actualizados',
        'criterios': [
            'El usuario puede editar su nombre completo',
            'El usuario puede editar su número de teléfono',
            'El usuario puede editar su biografía (máximo 500 caracteres)',
            'Se debe mostrar un contador de caracteres en tiempo real',
            'Se deben validar los campos antes de guardar',
            'Se debe mostrar confirmación al guardar cambios',
            'La interfaz se actualiza inmediatamente tras guardar'
        ],
        'prioridad': 'Alta',
        'notas': 'El email no se puede cambiar por seguridad. Endpoint: PUT /api/auth/profile'
    },
    {
        'id': 'RF-AU-06',
        'nombre': 'Cambiar Foto de Perfil',
        'modulo': 'Autenticación y Usuarios',
        'descripcion': 'Los usuarios pueden subir o cambiar su foto de perfil.',
        'como': 'Usuario autenticado',
        'quiero': 'Cambiar mi foto de perfil',
        'para': 'Personalizar mi cuenta',
        'criterios': [
            'Se acepta formato JPG, PNG, WebP',
            'El tamaño máximo es 5MB',
            'Se muestra preview antes de guardar',
            'La foto se almacena en el servidor',
            'La foto anterior se reemplaza al subir una nueva'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: POST /api/users/profile/photo con FormData'
    },
    {
        'id': 'RF-AU-07',
        'nombre': 'Cambiar Contraseña',
        'modulo': 'Autenticación y Usuarios',
        'descripcion': 'Los usuarios pueden cambiar su contraseña de acceso.',
        'como': 'Usuario autenticado',
        'quiero': 'Cambiar mi contraseña',
        'para': 'Mantener mi cuenta segura',
        'criterios': [
            'Se debe solicitar la contraseña actual',
            'Se debe solicitar la nueva contraseña (mínimo 6 caracteres)',
            'Se debe solicitar confirmación de la nueva contraseña',
            'Se valida que ambas contraseñas nuevas coincidan',
            'Se encripta la nueva contraseña con bcrypt',
            'Se muestra mensaje de éxito o error'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: PUT /api/users/change-password'
    },
    {
        'id': 'RF-AU-08',
        'nombre': 'Plan Premium',
        'modulo': 'Autenticación y Usuarios',
        'descripcion': 'Los usuarios pueden actualizar su cuenta a plan premium.',
        'como': 'Usuario autenticado',
        'quiero': 'Actualizar mi plan a premium',
        'para': 'Acceder a funcionalidades exclusivas',
        'criterios': [
            'Se muestra botón para actualizar a premium',
            'El cambio de plan se registra en la base de datos',
            'Se muestra badge visual de usuario premium',
            'Se actualiza el contexto del usuario inmediatamente'
        ],
        'prioridad': 'Baja',
        'notas': 'Endpoint: POST /api/users/upgrade-premium (gratis para proyecto académico)'
    },

    # MÓDULO 2: GESTIÓN DE RECETAS (12)
    {
        'id': 'RF-RE-01',
        'nombre': 'Ver Listado de Recetas',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'El sistema muestra un catálogo de recetas disponibles en formato grid responsive.',
        'como': 'Usuario (autenticado o no)',
        'quiero': 'Ver el catálogo de recetas',
        'para': 'Explorar las recetas disponibles',
        'criterios': [
            'Se muestra grid responsive (3 columnas en desktop, 2 en tablet, 1 en móvil)',
            'Cada tarjeta muestra: foto, nombre, categoría, dificultad, tiempo, rating',
            'Se implementa paginación de 12 recetas por página',
            'Las imágenes se cargan de forma optimizada',
            'Al hacer clic en una receta, se redirige al detalle'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/recipes. Página: /recipes'
    },
    {
        'id': 'RF-RE-02',
        'nombre': 'Buscar Recetas',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios pueden buscar recetas por nombre o ingredientes.',
        'como': 'Usuario',
        'quiero': 'Buscar recetas',
        'para': 'Encontrar recetas específicas rápidamente',
        'criterios': [
            'Hay una barra de búsqueda en el header',
            'La búsqueda funciona en tiempo real (debounce 300ms)',
            'Se busca en nombre de receta',
            'Se busca en ingredientes',
            'Los resultados se destacan visualmente',
            'Se actualiza el listado automáticamente'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/recipes?search=query'
    },
    {
        'id': 'RF-RE-03',
        'nombre': 'Filtrar Recetas por Categoría',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios pueden filtrar recetas por categoría.',
        'como': 'Usuario',
        'quiero': 'Filtrar por categoría',
        'para': 'Ver solo recetas de un tipo específico',
        'criterios': [
            'Hay filtros visuales con iconos de categorías',
            'Categorías: Tortas, Galletas, Postres Fríos, Chocolates, Tartas, Cupcakes, Otros',
            'Se muestra conteo de recetas por categoría',
            'El listado se actualiza automáticamente al filtrar'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/recipes?categoria=Tortas'
    },
    {
        'id': 'RF-RE-04',
        'nombre': 'Filtrar por Dificultad',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios pueden filtrar recetas por nivel de dificultad.',
        'como': 'Usuario',
        'quiero': 'Filtrar por dificultad',
        'para': 'Ver recetas acordes a mi nivel',
        'criterios': [
            'Opciones de filtro: Fácil, Intermedio, Difícil',
            'El filtro es combinable con otros filtros',
            'Indicador visual del nivel de dificultad',
            'Actualización automática del listado'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: GET /api/recipes?dificultad=Facil'
    },
    {
        'id': 'RF-RE-05',
        'nombre': 'Ver Detalle de Receta',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios pueden ver el detalle completo de una receta.',
        'como': 'Usuario',
        'quiero': 'Ver el detalle completo de una receta',
        'para': 'Conocer todos los ingredientes y pasos de preparación',
        'criterios': [
            'Se muestra información completa: nombre, descripción, categoría, dificultad, tiempo, porciones',
            'Lista de ingredientes con cantidades y unidades',
            'Pasos de preparación numerados',
            'Foto principal de la receta',
            'Nombre del autor',
            'Fecha de publicación',
            'Calificación promedio y total de valoraciones',
            'Botón para agregar a favoritos',
            'Sección de comentarios'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/recipes/:id. Página: /recipes/:id'
    },
    {
        'id': 'RF-RE-06',
        'nombre': 'Crear Receta (Wizard Multi-paso)',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios autenticados pueden crear recetas mediante un wizard de 4 pasos.',
        'como': 'Usuario autenticado',
        'quiero': 'Crear una nueva receta',
        'para': 'Compartir mis recetas con la comunidad',
        'criterios': [
            'Paso 1/4: Información básica (nombre, descripción, categoría, dificultad, tiempo, porciones)',
            'Paso 2/4: Ingredientes (agregar/editar/eliminar con cantidad y unidad)',
            'Paso 3/4: Pasos de preparación (numeración automática, reordenar)',
            'Paso 4/4: Foto principal (upload con preview)',
            'Navegación entre pasos con validación',
            'Preview antes de publicar',
            'Guardado exitoso con mensaje de confirmación'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: POST /api/recipes. Página: /user/create-recipe'
    },
    {
        'id': 'RF-RE-07',
        'nombre': 'Editar Receta Propia',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios pueden editar sus propias recetas.',
        'como': 'Usuario autenticado (autor)',
        'quiero': 'Editar mi receta',
        'para': 'Corregir o actualizar la información',
        'criterios': [
            'Solo el autor puede editar su receta',
            'Wizard igual al de creación',
            'Pre-carga de datos existentes',
            'Actualización de todos los campos',
            'Validación de permisos (solo autor o admin)'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: PUT /api/recipes/:id. Página: /user/edit-recipe/:id'
    },
    {
        'id': 'RF-RE-08',
        'nombre': 'Eliminar Receta Propia',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios pueden eliminar sus propias recetas.',
        'como': 'Usuario autenticado (autor)',
        'quiero': 'Eliminar mi receta',
        'para': 'Remover recetas que ya no quiero compartir',
        'criterios': [
            'Solo el autor puede eliminar su receta',
            'Modal de confirmación con advertencia',
            'Eliminación en cascada (ingredientes, pasos, comentarios)',
            'Mensaje de confirmación exitosa'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: DELETE /api/recipes/:id'
    },
    {
        'id': 'RF-RE-09',
        'nombre': 'Ver Mis Recetas',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios pueden ver todas sus recetas creadas.',
        'como': 'Usuario autenticado',
        'quiero': 'Ver todas mis recetas',
        'para': 'Gestionar mis publicaciones',
        'criterios': [
            'Listado de recetas del usuario autenticado',
            'Estados visibles: publicada, borrador, archivada',
            'Acciones rápidas: editar, eliminar, ver',
            'Ordenamiento por fecha de creación'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/recipes/my-recipes. Página: /user/my-recipes'
    },
    {
        'id': 'RF-RE-10',
        'nombre': 'Imprimir Receta',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Los usuarios pueden imprimir recetas en formato optimizado.',
        'como': 'Usuario',
        'quiero': 'Imprimir una receta',
        'para': 'Tenerla físicamente mientras cocino',
        'criterios': [
            'Vista de impresión limpia sin elementos de navegación',
            'Optimizada para papel tamaño A4',
            'Incluye: foto, ingredientes y pasos',
            'Botón de imprimir visible en el detalle de receta'
        ],
        'prioridad': 'Baja',
        'notas': 'Frontend: window.print(). Componente: PrintRecipe'
    },
    {
        'id': 'RF-RE-11',
        'nombre': 'Recetas Destacadas en Home',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'La página principal muestra recetas destacadas.',
        'como': 'Usuario',
        'quiero': 'Ver recetas destacadas al entrar',
        'para': 'Descubrir recetas populares rápidamente',
        'criterios': [
            'Carrusel de recetas destacadas en la página principal',
            'Ordenadas por calificación promedio',
            'Máximo 6 recetas destacadas',
            'Auto-scroll cada 5 segundos',
            'Navegación manual con flechas'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/recipes/featured'
    },
    {
        'id': 'RF-RE-12',
        'nombre': 'Recetas Recientes',
        'modulo': 'Gestión de Recetas',
        'descripcion': 'Mostrar las recetas publicadas recientemente.',
        'como': 'Usuario',
        'quiero': 'Ver las recetas más recientes',
        'para': 'Descubrir contenido nuevo',
        'criterios': [
            'Ordenadas por fecha de creación descendente',
            'Mostrar las últimas 12 recetas',
            'Grid responsive',
            'Link para ver todas las recetas'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: GET /api/recipes?sort=recent&limit=12'
    },

    # MÓDULO 3: FAVORITOS Y COLECCIONES (4)
    {
        'id': 'RF-FA-01',
        'nombre': 'Agregar a Favoritos',
        'modulo': 'Favoritos y Colecciones',
        'descripcion': 'Los usuarios pueden marcar recetas como favoritas.',
        'como': 'Usuario autenticado',
        'quiero': 'Marcar recetas como favoritas',
        'para': 'Acceder rápidamente a mis recetas preferidas',
        'criterios': [
            'Botón de corazón visible en cada tarjeta de receta',
            'Toggle on/off (agregar/quitar favorito)',
            'Feedback visual inmediato al hacer clic',
            'Persistencia en base de datos',
            'Validación de usuario autenticado'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: POST /api/users/favorites/:recipeId'
    },
    {
        'id': 'RF-FA-02',
        'nombre': 'Ver Favoritos',
        'modulo': 'Favoritos y Colecciones',
        'descripcion': 'Los usuarios pueden ver todas sus recetas favoritas.',
        'como': 'Usuario autenticado',
        'quiero': 'Ver todas mis recetas favoritas',
        'para': 'Acceder rápidamente a mis recetas guardadas',
        'criterios': [
            'Listado de recetas marcadas como favoritas',
            'Ordenadas por fecha de guardado descendente',
            'Grid responsive',
            'Botón para quitar de favoritos',
            'Contador total de favoritos'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/users/favorites. Página: /user/favorites'
    },
    {
        'id': 'RF-FA-03',
        'nombre': 'Crear Colección',
        'modulo': 'Favoritos y Colecciones',
        'descripcion': 'Los usuarios pueden crear colecciones personalizadas de recetas.',
        'como': 'Usuario autenticado',
        'quiero': 'Crear colecciones personalizadas',
        'para': 'Organizar mis recetas por temas',
        'criterios': [
            'Nombre de colección obligatorio',
            'Descripción opcional',
            'Modal de creación',
            'Validación de nombre único por usuario',
            'Confirmación de creación exitosa'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: POST /api/users/collections'
    },
    {
        'id': 'RF-FA-04',
        'nombre': 'Agregar Receta a Colección',
        'modulo': 'Favoritos y Colecciones',
        'descripcion': 'Los usuarios pueden agregar recetas a sus colecciones.',
        'como': 'Usuario autenticado',
        'quiero': 'Agregar recetas a mis colecciones',
        'para': 'Organizar mis recetas guardadas por categorías personales',
        'criterios': [
            'Selector de colección disponible',
            'Una receta puede estar en múltiples colecciones',
            'Validación de duplicados en la misma colección',
            'Feedback visual de éxito'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: POST /api/users/collections/:id/recipes/:recipeId'
    },

    # MÓDULO 4: VALORACIONES Y COMENTARIOS (3)
    {
        'id': 'RF-VC-01',
        'nombre': 'Valorar Receta',
        'modulo': 'Valoraciones y Comentarios',
        'descripcion': 'Los usuarios pueden calificar recetas con estrellas del 1 al 5.',
        'como': 'Usuario autenticado',
        'quiero': 'Calificar recetas',
        'para': 'Expresar mi opinión sobre la receta',
        'criterios': [
            'Sistema de 5 estrellas',
            'Un usuario solo puede valorar una vez por receta',
            'Actualización del promedio en tiempo real',
            'Contador de total de valoraciones',
            'Validación de usuario autenticado'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: POST /api/recipes/:id/rate'
    },
    {
        'id': 'RF-VC-02',
        'nombre': 'Comentar Receta',
        'modulo': 'Valoraciones y Comentarios',
        'descripcion': 'Los usuarios pueden dejar comentarios en las recetas.',
        'como': 'Usuario autenticado',
        'quiero': 'Comentar en recetas',
        'para': 'Compartir mi experiencia o hacer preguntas',
        'criterios': [
            'Caja de texto para escribir comentario',
            'Mínimo 10 caracteres',
            'Máximo 500 caracteres',
            'Mostrar nombre y foto del usuario autor',
            'Mostrar fecha del comentario',
            'Ordenados por más recientes primero'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: POST /api/recipes/:id/comments'
    },
    {
        'id': 'RF-VC-03',
        'nombre': 'Eliminar Comentario Propio',
        'modulo': 'Valoraciones y Comentarios',
        'descripcion': 'Los usuarios pueden eliminar sus propios comentarios.',
        'como': 'Usuario autenticado (autor del comentario)',
        'quiero': 'Eliminar mis comentarios',
        'para': 'Remover comentarios que ya no quiero mantener',
        'criterios': [
            'Solo el autor puede eliminar su propio comentario',
            'Modal de confirmación',
            'Actualización inmediata de la lista',
            'Mensaje de confirmación exitosa'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: DELETE /api/comments/:id'
    },

    # MÓDULO 5: PANEL DE ADMINISTRACIÓN (8)
    {
        'id': 'RF-AD-01',
        'nombre': 'Dashboard de Administración',
        'modulo': 'Panel de Administración',
        'descripcion': 'Panel con estadísticas y acceso rápido a funciones de administración.',
        'como': 'Administrador',
        'quiero': 'Ver un dashboard con estadísticas',
        'para': 'Monitorear el estado general del sistema',
        'criterios': [
            'Acceso exclusivo para usuarios con rol admin',
            'Tarjetas con métricas: total usuarios, total recetas, comentarios, categorías',
            'Gráficos de actividad',
            'Accesos rápidos a gestiones',
            'Diseño responsive'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/admin/stats. Página: /admin'
    },
    {
        'id': 'RF-AD-02',
        'nombre': 'Gestión de Usuarios',
        'modulo': 'Panel de Administración',
        'descripcion': 'Los administradores pueden gestionar usuarios del sistema.',
        'como': 'Administrador',
        'quiero': 'Gestionar usuarios',
        'para': 'Administrar roles y eliminar cuentas si es necesario',
        'criterios': [
            'Listado de todos los usuarios',
            'Búsqueda por nombre o email',
            'Paginación',
            'Ver datos: nombre, email, rol, fecha de registro',
            'Cambiar rol (usuario ↔ admin)',
            'Eliminar usuarios'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/admin/users. Página: /admin/users'
    },
    {
        'id': 'RF-AD-03',
        'nombre': 'Cambiar Rol de Usuario',
        'modulo': 'Panel de Administración',
        'descripcion': 'Los administradores pueden cambiar el rol de los usuarios.',
        'como': 'Administrador',
        'quiero': 'Cambiar el rol de usuarios',
        'para': 'Promover usuarios a administradores o viceversa',
        'criterios': [
            'Dropdown con opciones: usuario, admin',
            'Actualización inmediata en base de datos',
            'Confirmación visual del cambio',
            'El admin no puede cambiar su propio rol'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: PUT /api/admin/users/:id/role'
    },
    {
        'id': 'RF-AD-04',
        'nombre': 'Eliminar Usuario',
        'modulo': 'Panel de Administración',
        'descripcion': 'Los administradores pueden eliminar usuarios del sistema.',
        'como': 'Administrador',
        'quiero': 'Eliminar usuarios',
        'para': 'Remover cuentas inactivas o infractoras',
        'criterios': [
            'Modal de confirmación con advertencia clara',
            'Eliminación en cascada (recetas, comentarios, favoritos del usuario)',
            'El admin no puede eliminarse a sí mismo',
            'Mensaje de confirmación exitosa'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: DELETE /api/admin/users/:id'
    },
    {
        'id': 'RF-AD-05',
        'nombre': 'Gestión de Recetas',
        'modulo': 'Panel de Administración',
        'descripcion': 'Los administradores pueden gestionar todas las recetas del sistema.',
        'como': 'Administrador',
        'quiero': 'Gestionar todas las recetas',
        'para': 'Moderar contenido inapropiado o duplicado',
        'criterios': [
            'Listado de todas las recetas del sistema',
            'Búsqueda y filtros',
            'Ver autor, fecha, estado de cada receta',
            'Editar cualquier receta',
            'Eliminar recetas',
            'Cambiar estado (publicada/archivada)'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/admin/recipes. Página: /admin/recipes'
    },
    {
        'id': 'RF-AD-06',
        'nombre': 'Moderación de Comentarios',
        'modulo': 'Panel de Administración',
        'descripcion': 'Los administradores pueden moderar comentarios del sistema.',
        'como': 'Administrador',
        'quiero': 'Moderar comentarios',
        'para': 'Eliminar comentarios inapropiados u ofensivos',
        'criterios': [
            'Listado de todos los comentarios',
            'Ver autor, receta asociada, fecha',
            'Eliminar comentarios inapropiados',
            'Ordenamiento por más recientes',
            'Confirmación antes de eliminar'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: GET /api/admin/comments. Página: /admin/comments'
    },
    {
        'id': 'RF-AD-07',
        'nombre': 'Gestión de Categorías',
        'modulo': 'Panel de Administración',
        'descripcion': 'Los administradores pueden gestionar categorías de recetas.',
        'como': 'Administrador',
        'quiero': 'Gestionar categorías',
        'para': 'Agregar, editar o eliminar categorías de recetas',
        'criterios': [
            'Listado de categorías con conteo de recetas',
            'Crear nueva categoría (nombre, icono, color)',
            'Editar categoría existente',
            'Eliminar categoría (solo si no tiene recetas asociadas)',
            'Validación de nombre único'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoints: GET/POST/PUT/DELETE /api/admin/categories. Página: /admin/categories'
    },
    {
        'id': 'RF-AD-08',
        'nombre': 'Botón de Retroceso en Páginas Admin',
        'modulo': 'Panel de Administración',
        'descripcion': 'Todas las páginas de administración tienen botón para volver al dashboard.',
        'como': 'Administrador',
        'quiero': 'Volver fácilmente al dashboard',
        'para': 'Navegar cómodamente entre las secciones de admin',
        'criterios': [
            'BackButton visible en header de cada página admin',
            'Navegación directa a /admin',
            'Icono de flecha claramente visible',
            'Responsive en todos los dispositivos'
        ],
        'prioridad': 'Baja',
        'notas': 'Frontend: useNavigate(-1) o Link to="/admin"'
    },

    # MÓDULO 6: CATEGORÍAS (3)
    {
        'id': 'RF-CAT-01',
        'nombre': 'Listar Categorías',
        'modulo': 'Categorías',
        'descripcion': 'El sistema permite visualizar todas las categorías de recetas disponibles.',
        'como': 'Usuario',
        'quiero': 'Ver todas las categorías',
        'para': 'Conocer los tipos de recetas disponibles',
        'criterios': [
            'Listado completo de categorías',
            'Información mostrada: nombre, icono, color',
            'Contador de recetas por categoría',
            'Ordenamiento alfabético'
        ],
        'prioridad': 'Alta',
        'notas': 'Endpoint: GET /api/categories'
    },
    {
        'id': 'RF-CAT-02',
        'nombre': 'Crear Categoría (Admin)',
        'modulo': 'Categorías',
        'descripcion': 'Los administradores pueden crear nuevas categorías de recetas.',
        'como': 'Administrador',
        'quiero': 'Crear nuevas categorías',
        'para': 'Expandir las opciones de clasificación de recetas',
        'criterios': [
            'Formulario con campos: nombre, icono, color',
            'Validación de nombre único',
            'Preview del icono seleccionado',
            'Mensaje de confirmación al crear'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: POST /api/admin/categories'
    },
    {
        'id': 'RF-CAT-03',
        'nombre': 'Editar Categoría (Admin)',
        'modulo': 'Categorías',
        'descripcion': 'Los administradores pueden editar categorías existentes.',
        'como': 'Administrador',
        'quiero': 'Editar categorías',
        'para': 'Actualizar información de categorías existentes',
        'criterios': [
            'Modificar nombre, icono o color',
            'Actualización en tiempo real',
            'No afecta recetas existentes de la categoría',
            'Mensaje de confirmación'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: PUT /api/admin/categories/:id'
    },

    # MÓDULO 7: NOTIFICACIONES (2)
    {
        'id': 'RF-NOT-01',
        'nombre': 'Campana de Notificaciones',
        'modulo': 'Notificaciones',
        'descripcion': 'Los usuarios ven un ícono de campana con contador de notificaciones no leídas.',
        'como': 'Usuario autenticado',
        'quiero': 'Ver mis notificaciones',
        'para': 'Estar informado de actividad en mis recetas',
        'criterios': [
            'Ícono de campana visible en navbar',
            'Badge con contador de notificaciones no leídas',
            'Dropdown con listado de notificaciones al hacer clic',
            'Actualización en tiempo real',
            'Máximo 10 notificaciones visibles'
        ],
        'prioridad': 'Media',
        'notas': 'Endpoint: GET /api/notifications. Componente: NotificationBell'
    },
    {
        'id': 'RF-NOT-02',
        'nombre': 'Marcar Notificación como Leída',
        'modulo': 'Notificaciones',
        'descripcion': 'Los usuarios pueden marcar notificaciones como leídas.',
        'como': 'Usuario autenticado',
        'quiero': 'Marcar notificaciones como leídas',
        'para': 'Mantener organizada mi bandeja de notificaciones',
        'criterios': [
            'Click en notificación la marca como leída automáticamente',
            'Actualización del contador de no leídas',
            'Cambio de estilo visual (gris para leídas)',
            'Persistencia en base de datos'
        ],
        'prioridad': 'Baja',
        'notas': 'Endpoint: PUT /api/notifications/:id/read'
    }
]

# Generar tabla para cada requerimiento
for idx, req in enumerate(requerimientos, 1):
    # Crear tabla de 7 filas x 2 columnas
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    # Configurar ancho de columnas
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)

    # FILA 1: ID y Nombre
    cell_header = table.rows[0].cells[0]
    cell_header.merge(table.rows[0].cells[1])
    set_cell_background(cell_header, 'E8E8E8')  # Gris claro
    p = cell_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{req['id']}  {req['nombre']}")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    # FILA 2: Módulo
    cell_label = table.rows[1].cells[0]
    set_cell_background(cell_label, 'F5F5F5')  # Gris muy claro
    p = cell_label.paragraphs[0]
    run = p.add_run('Módulo:')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[1].cells[1]
    p = cell_value.paragraphs[0]
    run = p.add_run(req['modulo'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # FILA 3: Descripción
    cell_label = table.rows[2].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('DESCRIPCIÓN')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[2].cells[1]
    p = cell_value.paragraphs[0]
    run = p.add_run(req['descripcion'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # FILA 4: Como, Quiero, Para
    cell_label = table.rows[3].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('COMO, QUIERO, PARA')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[3].cells[1]
    p = cell_value.paragraphs[0]

    run = p.add_run('Como: ')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run = p.add_run(req['como'] + '\n')
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    run = p.add_run('Quiero: ')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run = p.add_run(req['quiero'] + '\n')
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    run = p.add_run('Para: ')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run = p.add_run(req['para'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # FILA 5: Criterios de Aceptación
    cell_label = table.rows[4].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('CRITERIOS DE ACEPTACIÓN')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[4].cells[1]
    p = cell_value.paragraphs[0]
    for i, criterio in enumerate(req['criterios']):
        if i > 0:
            p = cell_value.add_paragraph()
        run = p.add_run(f'• {criterio}')
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    # FILA 6: Prioridad
    cell_label = table.rows[5].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('PRIORIDAD')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[5].cells[1]
    p = cell_value.paragraphs[0]
    run = p.add_run(req['prioridad'])
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # FILA 7: Notas Adicionales
    cell_label = table.rows[6].cells[0]
    set_cell_background(cell_label, 'F5F5F5')
    p = cell_label.paragraphs[0]
    run = p.add_run('NOTAS ADICIONALES')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    cell_value = table.rows[6].cells[1]
    p = cell_value.paragraphs[0]
    run = p.add_run(req['notas'])
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.italic = True

    # Espacio entre tablas
    doc.add_paragraph()

    # Salto de página cada 2 requerimientos para mejor visualización
    if idx % 2 == 0 and idx < len(requerimientos):
        doc.add_page_break()

# Guardar documento
path = 'C:/xampp/htdocs/ProSweetBites/appnueva/SweetBites_Requerimientos_40_Completos.docx'
doc.save(path)

print(f"\n[OK] Documento generado exitosamente!")
print(f"Ubicación: {path}")
print(f"Total de requerimientos: {len(requerimientos)}")
print("Formato: Arial 12, sin colores, tablas individuales limpias")
